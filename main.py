import os
import docx
from datetime import date, datetime
from dateutil import parser
from spire import doc as sp_doc
from spire.doc import Document, FileFormat
# from docx2pdf import convert --FOR WINDOWS

os.environ["DOTNET_SYSTEM_GLOBALIZATION_INVARIANT"] = "true"

pdf_doc = Document()




# --Fields--

# address   -> default = ["E-414, Ganesh Glory 11, Jagatpur
#                    Rd, near BSNL Office Off. S.G Highway,
#                    Ahmedabad 382470"]
# contract_date
# address
# auth_per
# designation
# eff_date
# client_company
# client_address
# client_name
# client_designation
# gender

# --keep in mind--
# location
# honorifics  ->  (Mr. / Mrs.)


# --Declaration--
class main:

  dates = ["contract_date", "eff_date"]

  capitalize = [
      "auth_per",
      "address",
      "client_company",
      "client_address",
      "client_name",
  ]

  # ------FOR Console input------
  # ask = {
  #     "contract_date": "Enter Date of Contract(dd-mm-yyyy):",
  #     "address": "Enter address of company:",
  #     "auth_per": "Enter authorized person:",
  #     "designation": "Enter Authorized person's Designation:",
  #     "eff_date": "Enter effect date(dd-mm-yyyy):",
  #     "client_company": "Enter Client company:",
  #     "client_address": "Enter client Company address:",
  #     "client_name": "Enter Client Name:",
  #     "client_designation": "Enter client's Designation (client company):",
  #     "gender": "Gender(M/F):",
  # }

  placeholder = {
      "contract_date": "".strip(),
      "address": "",
      "auth_per": "",
      "designation": "",
      "eff_date": "".strip(),
      "client_company": "",
      "client_address": "",
      "client_name": "",
      "client_designation": "",
      "gender": ""
  }

  default = {
    "eff_date" : datetime.now().strftime("%m-%d-%Y"),
    "contract_date": datetime.now().strftime("%m-%d-%Y"),
    "address": "E-414, Ganesh Glory 11, Jagatpur Rd, near BSNL Office Off. S.G Highway, Ahmedabad 38247",
  }

  # --Getting Input--
  def get_input(self, placeholder):
    for key, val in placeholder.items():
      # placeholder[key] = input(f"{val}") # --for getting values from console
      if (placeholder[key] == ""):
        if key in self.default:
          placeholder[key] = self.default[key]
        else:
          placeholder[key] = "!!Empty Field!!"


      if key == "gender":
        if placeholder[key].lower() == "m" or placeholder[key].lower() == "male" :
          placeholder[key] = "Mr."
        elif placeholder[key].lower() == "f" or placeholder[key].lower() == "female":
          placeholder[key] = "Mrs."

      # --Capitalize--
      if key in self.capitalize:
        placeholder[key] = placeholder[key].capitalize()
      # --Format Date--
      try:
        if key in self.dates:
          placeholder[key] = parser.parse(
              placeholder[key]).strftime("%b %d, %Y")
      except Exception:
        print("\nWrong input For Date")

    # --Create Docx--
    doc = docx.Document("Formatted_document.docx")
    file_name = "NDA_" + placeholder["client_company"]

    for para in doc.paragraphs:
      for run in para.runs:
        txt = run.text
        for key, val in placeholder.items():
          if txt.find(key) == True:
            run.clear()
            run.add_text(val)
            run.italic = False
            break

    
#   --Save And Convert--
    doc.save(f"Updated/{file_name}.docx")
    # convert(f"{file_name}.docx", f"{file_name}.pdf") --FOR WINDOWS(docx2pdf)
    pdf_doc.LoadFromFile(f"Updated/{file_name}.docx", FileFormat.Docx)
    pdf_doc.SaveToFile(f"Updated/PDF/{file_name}.pdf", FileFormat.PDF)
    pdf_doc.Close()

